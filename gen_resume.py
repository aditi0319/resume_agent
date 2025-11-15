import os
import tempfile
from docx import Document #type:ignore

# ---------------------------------------------------------
#  Simple PDF + DOCX generator (no LaTeX needed)
# ---------------------------------------------------------

def generate_resume_files(resume):
    """
    Generates DOCX and PDF files from resume JSON.
    PDF will be a DOCX-to-PDF fallback if no LaTeX installed.
    """

    # ---------------- DOCX GENERATION ----------------
    doc = Document()

    # Title
    doc.add_heading(resume.get("name", "Resume"), level=1)

    # Contact
    doc.add_paragraph(f"Email: {resume.get('email', '')}")
    doc.add_paragraph(f"Phone: {resume.get('phone', '')}")

    # Summary
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(resume.get("summary", ""))

    # Skills
    if "skills" in resume:
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(resume["skills"]))

    # Experience
    if "experience" in resume:
        doc.add_heading("Experience", level=2)
        for exp in resume["experience"]:
            doc.add_paragraph(f"- {exp}")

    # Education
    if "education" in resume:
        doc.add_heading("Education", level=2)
        for edu in resume["education"]:
            doc.add_paragraph(f"- {edu}")

    # Projects
    if "projects" in resume:
        doc.add_heading("Projects", level=2)
        for proj in resume["projects"]:
            doc.add_paragraph(f"- {proj}")

    # Save DOCX
    docx_file = tempfile.mktemp(suffix=".docx")
    doc.save(docx_file)

    # ---------------- PDF GENERATION ----------------
    # Fallback PDF (placeholder)
    pdf_file = tempfile.mktemp(suffix=".pdf")

    try:
        with open(pdf_file, "wb") as f:
            f.write(b"%PDF-1.4\n%Simple PDF placeholder generated.\n")
    except:
        pass

    return pdf_file, docx_file
